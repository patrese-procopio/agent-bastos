"""
services/export_service.py
─────────────────────────────────────────────────────────────────────────────
Geração de laudos de análise fonográfica em TXT, PDF e DOCX.

Por que existe este módulo?
  - As funções _build_txt, _build_pdf e _build_docx viviam no api.py,
    misturadas com rotas HTTP e configuração global.
  - Aqui elas são funções puras: recebem um dict com os dados do laudo
    e devolvem bytes prontos para download. Zero dependência de FastAPI.
  - Isso permite testar a geração de cada formato com pytest sem subir
    o servidor — basta passar um dict com dados de exemplo.

Funções exportadas:
  build_txt(transcript: dict)  -> bytes   (UTF-8, plain text)
  build_pdf(transcript: dict)  -> bytes   (PDF via reportlab)
  build_docx(transcript: dict) -> bytes   (DOCX via python-docx)
"""

import io


# --- TXT ---------------------------------------------------------------------

def build_txt(t: dict) -> bytes:
    sep  = "=" * 60
    dash = "-" * 60
    lines = [
        sep, "SEAP/AM - AGENCIA DE INTELIGENCIA PENITENCIARIA",
        "LAUDO DE ANALISE FONOGRAFICA", sep,
        f"No: {t.get('laudo_number', 'N/A')}",
        f"Data: {t.get('date', 'N/A')}",
        f"Arquivo: {t.get('filename', 'N/A')}",
        f"Duracao: {t.get('duration', 'N/A')}",
        f"Nivel de Risco: {t.get('risk_level', 'N/A')}",
        f"Classificacao: {t.get('classification', 'N/A')}",
        dash, "INTERLOCUTORES:",
    ]
    for sp in t.get("speakers", []):
        lines.append(f"  {sp.get('id')} - {sp.get('label')} / {sp.get('role')}")
    lines += [dash, "TRANSCRICAO SEGMENTADA:"]
    for seg in t.get("segments", []):
        lines.append(f"[{seg.get('ts')}] {seg.get('speaker')}: {seg.get('text')}")
    lines += [dash, "RESUMO ANALITICO:", t.get("summary", "")]
    flags = t.get("red_flags", [])
    if flags:
        lines += [dash, "ALERTAS IDENTIFICADOS:"]
        for fl in flags:
            lines.append(f"  [{fl.get('id')}] {fl.get('title')}: {fl.get('text')}")
    lines += [sep, "Gerado pelo Agent Bastos - BASTOS-UNIT", sep]
    return "\n".join(lines).encode("utf-8")


# --- PDF ---------------------------------------------------------------------

def build_pdf(t: dict) -> bytes:
    """
    Gera laudo em PDF usando reportlab.
    Import lazy: reportlab so e carregado quando o usuario clica em exportar.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=2 * cm, bottomMargin=2 * cm,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
    )
    styles = getSampleStyleSheet()
    S = {
        "title": ParagraphStyle("T", parent=styles["Heading1"],
                                fontSize=13, alignment=TA_CENTER, spaceAfter=2),
        "sub":   ParagraphStyle("S", parent=styles["Normal"],
                                fontSize=9, alignment=TA_CENTER, spaceAfter=10,
                                textColor=colors.grey),
        "head":  ParagraphStyle("H", parent=styles["Heading2"],
                                fontSize=10, spaceBefore=10, spaceAfter=4,
                                textColor=colors.HexColor("#1E3A5F")),
        "body":  ParagraphStyle("B", parent=styles["Normal"], fontSize=9, leading=14),
        "mono":  ParagraphStyle("M", parent=styles["Normal"], fontSize=9, leading=13,
                                leftIndent=10, fontName="Courier"),
        "flag":  ParagraphStyle("F", parent=styles["Normal"], fontSize=9,
                                leftIndent=12, textColor=colors.HexColor("#DC2626")),
    }
    risk     = t.get("risk_level", "MEDIO")
    risk_hex = {"ALTO": "#DC2626", "MEDIO": "#D97706", "BAIXO": "#16A34A"}.get(risk, "#D97706")
    nav      = colors.HexColor("#1E3A5F")

    elements = [
        Paragraph("SEAP/AM - AGENCIA DE INTELIGENCIA PENITENCIARIA", S["title"]),
        Paragraph("LAUDO DE ANALISE FONOGRAFICA - CONFIDENCIAL", S["sub"]),
        HRFlowable(width="100%", thickness=1, color=nav), Spacer(1, 8),
    ]
    meta = [
        ["No do Laudo:", t.get("laudo_number", "N/A"), "Data:",    t.get("date", "N/A")],
        ["Arquivo:",     t.get("filename", "N/A"),      "Duracao:", t.get("duration", "N/A")],
        ["Classificacao:", t.get("classification", "N/A"), "Risco:",
         Paragraph(f'<font color="{risk_hex}"><b>{risk}</b></font>', S["body"])],
    ]
    mt = Table(meta, colWidths=[3.5 * cm, 7 * cm, 2.5 * cm, 4 * cm])
    mt.setStyle(TableStyle([
        ("FONTSIZE",      (0, 0), (-1, -1), 9),
        ("FONTNAME",      (0, 0), (0, -1),  "Helvetica-Bold"),
        ("FONTNAME",      (2, 0), (2, -1),  "Helvetica-Bold"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING",    (0, 0), (-1, -1), 3),
    ]))
    elements += [mt, Spacer(1, 10), HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey)]

    elements.append(Paragraph("INTERLOCUTORES", S["head"]))
    for sp in t.get("speakers", []):
        elements.append(Paragraph(
            f"<b>{sp.get('id')}</b> - {sp.get('label')} / {sp.get('role')}", S["body"]))

    elements.append(Paragraph("TRANSCRICAO SEGMENTADA", S["head"]))
    for seg in t.get("segments", []):
        elements.append(Paragraph(
            f"<b>[{seg.get('ts')}] {seg.get('speaker')}:</b> {seg.get('text')}", S["mono"]))

    elements.append(Paragraph("RESUMO ANALITICO", S["head"]))
    elements.append(Paragraph(t.get("summary", ""), S["body"]))

    flags = t.get("red_flags", [])
    if flags:
        elements.append(Paragraph("ALERTAS IDENTIFICADOS", S["head"]))
        for fl in flags:
            elements.append(Paragraph(
                f"<b>[{fl.get('id')}] {fl.get('title')}:</b> {fl.get('text')}", S["flag"]))

    elements += [
        Spacer(1, 16), HRFlowable(width="100%", thickness=1, color=nav),
        Paragraph("Documento gerado pelo sistema Agent Bastos - BASTOS-UNIT", S["sub"]),
    ]
    doc.build(elements)
    buf.seek(0)
    return buf.read()


# --- DOCX --------------------------------------------------------------------

def build_docx(t: dict) -> bytes:
    """
    Gera laudo em DOCX usando python-docx.
    Import lazy: python-docx so e carregado quando o usuario exporta em Word.
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    RISK_RGB = {
        "ALTO":  RGBColor(0xDC, 0x26, 0x26),
        "MEDIO": RGBColor(0xD9, 0x77, 0x06),
        "BAIXO": RGBColor(0x16, 0xA3, 0x4A),
    }
    risk     = t.get("risk_level", "MEDIO")
    risk_rgb = RISK_RGB.get(risk, RISK_RGB["MEDIO"])
    grey     = RGBColor(0x6B, 0x72, 0x80)
    red      = RGBColor(0xDC, 0x26, 0x26)

    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    h = doc.add_heading("SEAP/AM - AGENCIA DE INTELIGENCIA PENITENCIARIA", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("LAUDO DE ANALISE FONOGRAFICA - CONFIDENCIAL")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = grey
    sub.runs[0].font.size      = Pt(10)
    doc.add_paragraph()

    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    for i, (k1, v1, k2, v2) in enumerate([
        ("No do Laudo",    t.get("laudo_number", "N/A"), "Data",    t.get("date", "N/A")),
        ("Arquivo",        t.get("filename", "N/A"),      "Duracao", t.get("duration", "N/A")),
        ("Classificacao",  t.get("classification", "N/A"), "Risco",  risk),
    ]):
        row = table.rows[i]
        for j, (txt, bold) in enumerate([(k1, True), (v1, False), (k2, True), (v2, False)]):
            cell          = row.cells[j]
            cell.text     = txt
            run           = cell.paragraphs[0].runs[0]
            run.bold      = bold
            run.font.size = Pt(9)
            if i == 2 and j == 3:
                run.font.color.rgb = risk_rgb
    doc.add_paragraph()

    doc.add_heading("INTERLOCUTORES", 2)
    for sp in t.get("speakers", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(sp.get("id", "") + " - ").bold = True
        p.add_run(f"{sp.get('label', '')} / {sp.get('role', '')}")

    doc.add_heading("TRANSCRICAO SEGMENTADA", 2)
    for seg in t.get("segments", []):
        p  = doc.add_paragraph()
        r1 = p.add_run(f"[{seg.get('ts', '')}] {seg.get('speaker', '')}: ")
        r1.bold      = True
        r1.font.name = "Courier New"
        r1.font.size = Pt(9)
        r2           = p.add_run(seg.get("text", ""))
        r2.font.name = "Courier New"
        r2.font.size = Pt(9)

    doc.add_heading("RESUMO ANALITICO", 2)
    doc.add_paragraph(t.get("summary", ""))

    flags = t.get("red_flags", [])
    if flags:
        doc.add_heading("ALERTAS IDENTIFICADOS", 2)
        for fl in flags:
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(f"{fl.get('title', '')}: ")
            r.bold           = True
            r.font.color.rgb = red
            p.add_run(fl.get("text", ""))

    doc.add_paragraph()
    footer = doc.add_paragraph("Documento gerado pelo sistema Agent Bastos - BASTOS-UNIT")
    footer.alignment           = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.color.rgb = grey
    footer.runs[0].font.size      = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# --- RAE (Relatório Analítico de Extrato) ------------------------------------

def _esc(s) -> str:
    """Escapa caracteres sensíveis ao mini-XML do reportlab Paragraph."""
    return (str(s or "")
            .replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def build_rae_pdf(rae: dict) -> bytes:
    """
    Gera o RAE — Relatório Analítico de Extrato — em PDF timbrado.
    NÃO é um RELINT (produto final consolidado da agência); é o artefato
    analítico de nível de extrato, com validação humana.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable,
    )

    ex   = rae.get("extrato", {})
    nav  = colors.HexColor("#1E3A5F")
    gold = colors.HexColor("#B8860B")
    nivel = (rae.get("risk_nivel") or "MÉDIO").upper()
    risk_hex = {"ALTO": "#DC2626", "MÉDIO": "#D97706", "BAIXO": "#16A34A"}.get(nivel, "#D97706")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm,
                            leftMargin=2.2 * cm, rightMargin=2.2 * cm)
    styles = getSampleStyleSheet()
    S = {
        "title": ParagraphStyle("T", parent=styles["Heading1"], fontSize=13,
                                alignment=TA_CENTER, spaceAfter=2, textColor=nav),
        "sub":   ParagraphStyle("S", parent=styles["Normal"], fontSize=8.5,
                                alignment=TA_CENTER, spaceAfter=8, textColor=colors.grey),
        "head":  ParagraphStyle("H", parent=styles["Heading2"], fontSize=10.5,
                                spaceBefore=12, spaceAfter=4, textColor=nav),
        "body":  ParagraphStyle("B", parent=styles["Normal"], fontSize=9, leading=14),
        "small": ParagraphStyle("Sm", parent=styles["Normal"], fontSize=8, leading=11,
                                textColor=colors.grey),
        "flag":  ParagraphStyle("F", parent=styles["Normal"], fontSize=9, leading=13,
                                leftIndent=8, textColor=colors.HexColor("#B45309")),
    }

    el = [
        Paragraph("SEAP/AM &mdash; AG&Ecirc;NCIA DE INTELIG&Ecirc;NCIA PENITENCI&Aacute;RIA", S["title"]),
        Paragraph("RELAT&Oacute;RIO ANAL&Iacute;TICO DE EXTRATO (RAE) &middot; USO RESERVADO", S["sub"]),
        Paragraph("Artefato anal&iacute;tico de n&iacute;vel de extrato &mdash; n&atilde;o constitui RELINT. "
                  "Requer valida&ccedil;&atilde;o humana antes de consolida&ccedil;&atilde;o.", S["small"]),
        HRFlowable(width="100%", thickness=1, color=nav), Spacer(1, 6),
    ]

    meta = [
        ["Extrato:", _esc(ex.get("id")), "Data:", _esc(ex.get("data") or "—")],
        ["Unidade:", _esc(ex.get("unidade") or "—"), "Núcleo:", _esc(ex.get("nucleo") or "—")],
        ["Autor:", _esc(ex.get("autor") or "—"), "Classificação:", _esc(ex.get("classificacao") or "—")],
        ["Motor IA:", _esc(f"{ex.get('provedor') or '—'} / {ex.get('modelo') or '—'}"),
         "Risco:", Paragraph(f'<font color="{risk_hex}"><b>{_esc(nivel)} ({rae.get("risk_score")})</b></font>', S["body"])],
    ]
    mt = Table(meta, colWidths=[3 * cm, 7.6 * cm, 3 * cm, 3 * cm])
    mt.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4), ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("LINEBELOW", (0, 0), (-1, -2), 0.3, colors.whitesmoke),
    ]))
    el += [mt, Spacer(1, 8), HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey)]

    el.append(Paragraph("SÍNTESE DA AMEAÇA", S["head"]))
    el.append(Paragraph(_esc(rae.get("assunto_sintetizado") or "—"), S["body"]))

    just = rae.get("justificativa_risco") or "—"
    if rae.get("risco_forcado"):
        just += "  [Piso de risco aplicado automaticamente por palavra-crítica.]"
    el.append(Paragraph("AVALIAÇÃO DE RISCO", S["head"]))
    el.append(Paragraph(_esc(just), S["body"]))

    ents = rae.get("entidades", [])
    if ents:
        el.append(Paragraph(f"ENTIDADES-CHAVE ({len(ents)})", S["head"]))
        linhas = [["Tipo", "Rótulo", "Nome/Vulgo", "Papel no contexto", "Prov."]]
        for e in ents:
            nv = " / ".join([x for x in [e.get("nome"), e.get("vulgo")] if x]) or "—"
            prov = "✓" if e.get("evidencia_ok") else "⚠"
            linhas.append([
                _esc((e.get("tipo") or "").upper()),
                Paragraph(_esc(e.get("rotulo") or "—"), S["body"]),
                Paragraph(_esc(nv), S["body"]),
                Paragraph(_esc(e.get("papel") or "—"), S["body"]),
                prov,
            ])
        t = Table(linhas, colWidths=[2.2 * cm, 3.3 * cm, 3.6 * cm, 6 * cm, 1.2 * cm])
        t.setStyle(TableStyle([
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("BACKGROUND", (0, 0), (-1, 0), nav),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.3, colors.lightgrey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        el.append(t)

    cxs = rae.get("conexoes", [])
    if cxs:
        el.append(Paragraph(f"VÍNCULOS DETECTADOS ({len(cxs)})", S["head"]))
        for c in cxs:
            el.append(Paragraph(
                f"<b>{_esc(c.get('source'))}</b> &rarr; <i>{_esc(c.get('relation'))}</i> "
                f"&rarr; <b>{_esc(c.get('target'))}</b> &nbsp;(peso {_esc(c.get('weight'))})", S["body"]))

    jgs = rae.get("jargoes", [])
    if jgs:
        el.append(Paragraph(f"SINAIS FRACOS / JARGÕES ({len(jgs)})", S["head"]))
        for j in jgs:
            el.append(Paragraph(
                f"<b>«{_esc(j.get('termo'))}»</b> &asymp; {_esc(j.get('significado_provavel'))}", S["flag"]))

    tags = rae.get("tags", [])
    if tags:
        el.append(Paragraph("TAGS DE INDEXAÇÃO", S["head"]))
        el.append(Paragraph(_esc(" · ".join(tags)), S["small"]))

    el.append(Spacer(1, 10))
    el.append(Paragraph(
        f"Proveniência: {rae.get('evidencias_ok', 0)}/{rae.get('evidencias_total', 0)} "
        "itens com trecho-fonte conferido (✓). Itens com ⚠ exigem revisão manual.", S["small"]))

    el += [
        Spacer(1, 12), HRFlowable(width="100%", thickness=1, color=gold),
        Paragraph("Gerado pelo Agent Bastos — Módulo Extrato. Documento sob sigilo.", S["sub"]),
    ]
    doc.build(el)
    buf.seek(0)
    return buf.read()
