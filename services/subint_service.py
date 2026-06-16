# -*- coding: utf-8 -*-
"""
services/subint_service.py — SUBINT Automatizado (Missão 29)
─────────────────────────────────────────────────────────────────────────────
Gera Subsídios de Inteligência (SUBINT) automaticamente a partir dos dados
coletados e cruzados pelo Agent Bastos.

SUBINT ≠ RELINT:
  O SUBINT é um documento técnico preliminar que consolida os dados do sistema
  sobre uma entidade ou evento para SUBSIDIAR a elaboração do RELINT pelos
  analistas. Exportável em PDF e DOCX para cópia/cola no RELINT final.

FLUXO:
  1. Analista clica "Gerar SUBINT" no ORÁCULO
  2. Sistema coleta dados de todos os módulos sobre a entidade
  3. LLM (Claude Opus) sintetiza em documento estruturado
  4. Exporta para DOCX (editável pelo analista) e PDF (referência)
  5. Analista agrega conteúdo no RELINT manual

FONTES:
  • HITLs confirmados (human_loop_service)
  • Score de risco (risco_score_service)
  • Grafo de vínculos (modules.grafo)

LGPD:
  Documentos ficam em auth.db com registro de quem gerou e quando.
  Apenas analistas autenticados com módulo "admin" podem gerar/acessar.
"""

from __future__ import annotations

import io
import json
import logging
import os
import re
import unicodedata
import uuid
from datetime import datetime, timezone
from typing import Optional

logger = logging.getLogger("bastos.subint")

# ── DB ────────────────────────────────────────────────────────────────────────

def _db_path() -> str:
    return os.getenv(
        "AUTH_DB",
        os.path.join(os.path.dirname(__file__), "..", "data", "auth.db"),
    )


def _conn():
    import sqlite3
    con = sqlite3.connect(_db_path())
    con.row_factory = sqlite3.Row
    con.execute("PRAGMA journal_mode=WAL")
    return con


def _agora() -> str:
    return datetime.now(timezone.utc).isoformat()


def _init_db() -> None:
    with _conn() as con:
        con.executescript("""
            CREATE TABLE IF NOT EXISTS subint_docs (
                id           TEXT PRIMARY KEY,
                numero       TEXT NOT NULL,
                entidade     TEXT NOT NULL,
                origem       TEXT NOT NULL,
                operador     TEXT NOT NULL,
                texto        TEXT NOT NULL,
                pdf_bytes    BLOB,
                docx_bytes   BLOB,
                criado_em    TEXT NOT NULL
            );
            CREATE INDEX IF NOT EXISTS idx_subint_entidade
                ON subint_docs (entidade, criado_em DESC);
        """)


_init_db()


# ── Helpers ───────────────────────────────────────────────────────────────────

def _norm(txt: str) -> str:
    if not txt:
        return ""
    t = unicodedata.normalize("NFKD", str(txt))
    t = "".join(c for c in t if not unicodedata.combining(c))
    return re.sub(r"\s+", " ", t.strip().lower())


def _proximo_numero() -> str:
    """Gera número sequencial: 001/AIPEN/2026"""
    year = datetime.now().year
    with _conn() as con:
        n = con.execute(
            "SELECT COUNT(*) FROM subint_docs WHERE criado_em LIKE ?",
            (f"{year}%",),
        ).fetchone()[0]
    return f"{n+1:03d}/AIPEN/{year}"


# ── Coleta de dados ───────────────────────────────────────────────────────────

def _coletar_dados(entidade_nome: str, hitl_id: Optional[str] = None) -> dict:
    """
    Coleta todos os dados disponíveis sobre a entidade nos módulos do sistema.
    Fail-safe: cada fonte é independente — falha em uma não bloqueia as outras.
    """
    dados: dict = {
        "entidade":       entidade_nome,
        "score":          None,
        "historico_score": [],
        "nos_grafo":      [],
        "arestas":        [],
        "hitls":          [],
        "hits_hitl":      [],
        "hitl_descricao": "",
        "hitl_risco":     "ALTO",
        "hitl_summary":   "",
        "risco_nivel":    "DESCONHECIDO",
    }

    # ── Score de risco ────────────────────────────────────────────────────────
    try:
        from services.risco_score_service import _id_entidade, obter_score, historico_entidade
        eid   = _id_entidade(entidade_nome)
        score = obter_score(eid)
        if score:
            dados["score"]           = score
            dados["risco_nivel"]     = score.get("nivel", "DESCONHECIDO")
            dados["historico_score"] = historico_entidade(eid, limite=8)
    except Exception as exc:
        logger.warning("[subint] Score indisponível: %s", exc)

    # ── Grafo de vínculos ─────────────────────────────────────────────────────
    try:
        from modules.grafo import _conn as _gconn
        nome_norm = _norm(entidade_nome)
        with _gconn() as gcon:
            nos = gcon.execute(
                "SELECT id, tipo, rotulo, propriedades FROM nos "
                "WHERE lower(rotulo) LIKE ? LIMIT 5",
                (f"%{nome_norm}%",),
            ).fetchall()
            dados["nos_grafo"] = [dict(n) for n in nos]

            if nos:
                no_id   = nos[0]["id"]
                arestas = gcon.execute(
                    "SELECT rotulo, origem_id, destino_id FROM arestas "
                    "WHERE origem_id=? OR destino_id=? LIMIT 20",
                    (no_id, no_id),
                ).fetchall()
                # Enriquecer com rótulos dos nós conectados
                ids_extra = set()
                for a in arestas:
                    ids_extra.add(a["origem_id"])
                    ids_extra.add(a["destino_id"])
                ids_extra.discard(no_id)
                rotulos = {}
                for eid_g in ids_extra:
                    row = gcon.execute(
                        "SELECT rotulo FROM nos WHERE id=?", (eid_g,)
                    ).fetchone()
                    if row:
                        rotulos[eid_g] = row["rotulo"]
                dados["arestas"] = [
                    {
                        "rotulo":    a["rotulo"],
                        "de":        rotulos.get(a["origem_id"], a["origem_id"][:8]),
                        "para":      rotulos.get(a["destino_id"], a["destino_id"][:8]),
                    }
                    for a in arestas
                ]
    except Exception as exc:
        logger.warning("[subint] Grafo indisponível: %s", exc)

    # ── HITLs confirmados que mencionam a entidade ────────────────────────────
    try:
        from services.human_loop_service import listar_aprovacoes
        confirmados = listar_aprovacoes(status="confirmada", limite=50)
        nome_norm   = _norm(entidade_nome)
        relevantes  = []
        for h in confirmados:
            det = h.get("detalhes") or {}
            if isinstance(det, str):
                try:
                    det = json.loads(det)
                except Exception:
                    det = {}
            hits = det.get("hits", [])
            menciona = any(nome_norm in _norm(hit.get("nome", "")) for hit in hits)
            if menciona or nome_norm in _norm(h.get("descricao", "")):
                relevantes.append({
                    "id":        h.get("id", ""),
                    "descricao": h.get("descricao", ""),
                    "risco":     h.get("risco", ""),
                    "criado_em": h.get("criado_em", "")[:10],
                    "hits":      hits,
                    "summary":   det.get("summary", ""),
                })
        dados["hitls"] = relevantes[:8]
    except Exception as exc:
        logger.warning("[subint] HITLs indisponíveis: %s", exc)

    # ── HITL específico (se informado) ────────────────────────────────────────
    if hitl_id:
        try:
            from services.human_loop_service import buscar_aprovacao
            h = buscar_aprovacao(hitl_id)
            if h:
                det = h.get("detalhes") or {}
                if isinstance(det, str):
                    try:
                        det = json.loads(det)
                    except Exception:
                        det = {}
                dados["hits_hitl"]      = det.get("hits", [])
                dados["hitl_descricao"] = h.get("descricao", "")
                dados["hitl_risco"]     = h.get("risco", "ALTO")
                dados["hitl_summary"]   = det.get("summary", "")
        except Exception as exc:
            logger.warning("[subint] HITL específico indisponível: %s", exc)

    return dados


# ── Síntese via LLM (Claude) ──────────────────────────────────────────────────

def _sintetizar_via_llm(dados: dict, origem: str) -> str:
    """
    Chama Claude Opus para sintetizar os dados coletados em um SUBINT estruturado.
    O LLM recebe apenas os dados já coletados — não inventa informações.
    """
    import anthropic
    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

    entidade = dados["entidade"].upper()

    # Formata score
    score_bloco = "Sem registro de score no sistema."
    if dados.get("score"):
        s = dados["score"]
        score_bloco = (
            f"Score atual: {s['score_atual']}/100 — Nível {s['nivel']}\n"
            f"Último evento registrado: {s.get('ultimo_evento', 'N/A')[:10]}"
        )
        if dados.get("historico_score"):
            score_bloco += "\nHistórico recente:\n"
            for ev in dados["historico_score"][:5]:
                score_bloco += f"  • [{ev.get('criado_em','')[:10]}] {ev.get('tipo_evento','')} (+{ev.get('delta',0):.0f} pts)\n"

    # Formata HITLs
    hitls_bloco = "Nenhuma correlação confirmada no sistema."
    if dados.get("hitls"):
        hitls_bloco = ""
        for h in dados["hitls"][:5]:
            hitls_bloco += f"• [{h['risco']}] {h['descricao']} — {h['criado_em']}\n"
            if h.get("summary"):
                hitls_bloco += f"  Contexto: {h['summary']}\n"
            if h.get("hits"):
                nomes = [hh.get("nome","") for hh in h["hits"] if hh.get("nome")]
                if nomes:
                    hitls_bloco += f"  Co-mencionados: {', '.join(nomes[:5])}\n"

    # Formata hits do HITL específico
    hits_bloco = ""
    if dados.get("hits_hitl"):
        for h in dados["hits_hitl"]:
            hits_bloco += f"• {h.get('nome','?')} ({h.get('fonte','?')})"
            if h.get("detalhe"):
                hits_bloco += f": {h['detalhe']}"
            hits_bloco += "\n"
    if dados.get("hitl_summary"):
        hits_bloco += f"\nResumo da origem: {dados['hitl_summary']}\n"

    # Formata grafo
    grafo_bloco = "Sem vínculos mapeados no grafo de inteligência."
    if dados.get("arestas"):
        grafo_bloco = ""
        for a in dados["arestas"][:10]:
            grafo_bloco += f"• {a['de']} —[{a['rotulo']}]→ {a['para']}\n"

    prompt = f"""Você é um analista sênior de inteligência de segurança pública do Brasil.

Produza um SUBINT (Subsídio de Inteligência) técnico e objetivo sobre a entidade abaixo.
Use APENAS os dados fornecidos. NÃO invente informações, ocorrências ou vínculos não listados.
Escreva em português formal de inteligência de segurança pública.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
ENTIDADE: {entidade}
ORIGEM DO SUBINT: {origem}

SCORE DE RISCO:
{score_bloco}

CORRELAÇÕES CONFIRMADAS PELO ORÁCULO:
{hitls_bloco}

ENTIDADES E FONTES CO-MENCIONADAS:
{hits_bloco or "Nenhuma disponível."}

GRAFO DE VÍNCULOS:
{grafo_bloco}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

ESTRUTURA OBRIGATÓRIA (use exatamente estes títulos de seção):

1. IDENTIFICAÇÃO E CONTEXTO
[Apresente quem é a entidade, de onde aparecem no sistema, seu histórico no Agent Bastos]

2. CORRELAÇÕES IDENTIFICADAS
[Liste cada correlação confirmada com data, risco e entidades associadas]

3. REDE DE VÍNCULOS
[Descreva os vínculos do grafo de forma analítica — quem se conecta a quem e como]

4. AVALIAÇÃO DE RISCO
[Interprete o score, o nível de risco e o histórico de eventos]

5. CONSIDERAÇÕES FINAIS E DILIGÊNCIAS SUGERIDAS
[Conclusões analíticas e recomendações de acompanhamento para os analistas]

REGRAS:
- Use CAIXA ALTA para nomes de pessoas e organizações criminosas
- Seja preciso e conciso — cada seção máximo 3 parágrafos
- Se uma seção não tiver dados suficientes, escreva "Sem dados suficientes para análise nesta seção."
- NÃO inclua cabeçalho (data, número, origem) — ele será adicionado pelo sistema

Produza APENAS o corpo do documento com as 5 seções."""

    msg = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=2500,
        messages=[{"role": "user", "content": prompt}],
    )
    return msg.content[0].text


# ── Exportação DOCX ───────────────────────────────────────────────────────────

def _gerar_docx(numero: str, entidade: str, origem: str, texto: str, operador: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)

    NAVY = RGBColor(0x1A, 0x1A, 0x2E)
    GRAY = RGBColor(0x44, 0x44, 0x44)
    DIM  = RGBColor(0x88, 0x88, 0x88)

    def _add(text, size=10.5, bold=False, color=None, align=None, italic=False):
        p   = doc.add_paragraph()
        run = p.add_run(text)
        run.bold        = bold
        run.italic      = italic
        run.font.size   = Pt(size)
        run.font.color.rgb = color or GRAY
        if align:
            p.alignment = align
        p.paragraph_format.space_after = Pt(3)
        return p

    # Cabeçalho
    _add("SUBSÍDIO DE INTELIGÊNCIA — AGENT BASTOS", 14, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add(f"SUBINT Nº {numero}", 12, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    data_hoje = datetime.now().strftime("%d%b%Y").upper()
    for label, valor in [
        ("1. DATA:", data_hoje),
        ("2. ORIGEM:", origem),
        ("3. DIFUSÃO:", "INTERNA — USO EXCLUSIVO DA AIPEN"),
        ("4. ASSUNTO:", f"Informações sobre {entidade.upper()}"),
    ]:
        p  = doc.add_paragraph()
        r1 = p.add_run(f"{label} ")
        r1.bold      = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = NAVY
        r2 = p.add_run(valor)
        r2.font.size = Pt(11)
        r2.font.color.rgb = GRAY

    doc.add_paragraph()
    _add("─" * 90, size=7, color=DIM)
    doc.add_paragraph()

    # Corpo gerado pelo LLM
    for linha in texto.split("\n"):
        if not linha.strip():
            doc.add_paragraph()
            continue
        if re.match(r"^\d+\.\s+[A-ZÁÉÍÓÚÂÊÎÔÛÃÕÇ]", linha):
            _add(linha, size=11.5, bold=True, color=NAVY)
        elif linha.startswith(("- ", "• ")):
            p  = doc.add_paragraph(style="List Bullet" if "List Bullet" in [s.name for s in doc.styles] else "Normal")
            p.add_run(linha[2:]).font.size = Pt(10.5)
        else:
            _add(linha, size=10.5)

    doc.add_paragraph()
    _add("─" * 90, size=7, color=DIM)
    _add(
        '"É vedada a retransmissão parcial ou total deste documento/conhecimento '
        'sem a permissão desta AIPEN."',
        size=9, italic=True, color=DIM, align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add(
        f"Gerado automaticamente pelo AGENT BASTOS · Operador: {operador} · "
        f"{datetime.now().strftime('%d/%m/%Y %H:%M')}",
        size=8, color=DIM, align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Exportação PDF ────────────────────────────────────────────────────────────

def _gerar_pdf(numero: str, entidade: str, origem: str, texto: str, operador: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

    buf     = io.BytesIO()
    doc_pdf = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=2.5*cm, leftMargin=3*cm,
        topMargin=2.5*cm,   bottomMargin=2.5*cm,
        title=f"SUBINT {numero}",
    )

    styles = getSampleStyleSheet()
    NAVY   = colors.HexColor("#1A1A2E")
    GRAY   = colors.HexColor("#444444")
    LIGHT  = colors.HexColor("#888888")

    def S(name, **kw):
        return ParagraphStyle(name, parent=styles["Normal"], **kw)

    s_title = S("T", fontSize=14, fontName="Helvetica-Bold", textColor=NAVY, alignment=TA_CENTER, spaceAfter=4)
    s_sub   = S("S", fontSize=12, fontName="Helvetica-Bold", textColor=NAVY, alignment=TA_CENTER, spaceAfter=14)
    s_lbl   = S("L", fontSize=11, fontName="Helvetica-Bold", textColor=NAVY, spaceAfter=3)
    s_body  = S("B", fontSize=10, fontName="Helvetica",      textColor=GRAY, leading=15, alignment=TA_JUSTIFY, spaceAfter=5)
    s_head  = S("H", fontSize=11, fontName="Helvetica-Bold", textColor=NAVY, spaceBefore=10, spaceAfter=4)
    s_foot  = S("F", fontSize=8,  fontName="Helvetica-Oblique", textColor=LIGHT, alignment=TA_CENTER, spaceBefore=6)

    story   = []
    data_ok = datetime.now().strftime("%d%b%Y").upper()

    story += [
        Paragraph("SUBSÍDIO DE INTELIGÊNCIA — AGENT BASTOS", s_title),
        Paragraph(f"SUBINT Nº {numero}", s_sub),
        HRFlowable(width="100%", thickness=1.5, color=NAVY),
        Spacer(1, 0.3*cm),
    ]

    for label, valor in [
        ("1. DATA:", data_ok),
        ("2. ORIGEM:", origem),
        ("3. DIFUSÃO:", "INTERNA — USO EXCLUSIVO DA AIPEN"),
        ("4. ASSUNTO:", f"Informações sobre {entidade.upper()}"),
    ]:
        story.append(Paragraph(f"<b>{label}</b> {valor}", s_lbl))

    story += [
        Spacer(1, 0.4*cm),
        HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey),
        Spacer(1, 0.3*cm),
    ]

    for linha in texto.split("\n"):
        if not linha.strip():
            story.append(Spacer(1, 0.2*cm))
        elif re.match(r"^\d+\.\s+[A-ZÁÉÍÓÚÂÊÎÔÛÃÕÇ]", linha):
            story.append(Paragraph(linha, s_head))
        elif linha.startswith(("- ", "• ")):
            story.append(Paragraph(f"• {linha[2:]}", s_body))
        else:
            safe = linha.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, s_body))

    story += [
        Spacer(1, 0.5*cm),
        HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey),
        Paragraph(
            '"É vedada a retransmissão parcial ou total deste documento/conhecimento '
            'sem a permissão desta AIPEN."',
            s_foot,
        ),
        Paragraph(
            f"Gerado automaticamente pelo AGENT BASTOS · Operador: {operador} · "
            f"{datetime.now().strftime('%d/%m/%Y %H:%M')}",
            s_foot,
        ),
    ]

    doc_pdf.build(story)
    return buf.getvalue()


# ── Interface pública ─────────────────────────────────────────────────────────

def gerar_subint(
    entidade_nome: str,
    origem:        str,
    operador:      str,
    hitl_id:       Optional[str] = None,
) -> dict:
    """
    Gera um SUBINT completo para uma entidade ou evento.

    1. Coleta dados de todos os módulos do sistema
    2. Chama Claude Opus para síntese analítica
    3. Exporta para DOCX e PDF
    4. Persiste em auth.db

    Retorna: {"ok": True, "subint_id", "numero", "entidade", "criado_em"}
    """
    if not entidade_nome or not entidade_nome.strip():
        return {"ok": False, "motivo": "entidade_nome vazio"}

    subint_id = str(uuid.uuid4())
    numero    = _proximo_numero()

    try:
        logger.info("[subint] Gerando SUBINT %s para '%s' (operador=%s)", numero, entidade_nome, operador)

        dados      = _coletar_dados(entidade_nome, hitl_id)
        texto      = _sintetizar_via_llm(dados, origem)
        pdf_bytes  = _gerar_pdf(numero, entidade_nome, origem, texto, operador)
        docx_bytes = _gerar_docx(numero, entidade_nome, origem, texto, operador)
        agora      = _agora()

        with _conn() as con:
            con.execute(
                "INSERT INTO subint_docs "
                "(id, numero, entidade, origem, operador, texto, pdf_bytes, docx_bytes, criado_em) "
                "VALUES (?,?,?,?,?,?,?,?,?)",
                (subint_id, numero, entidade_nome, origem, operador,
                 texto, pdf_bytes, docx_bytes, agora),
            )

        logger.info("[subint] SUBINT %s gerado com sucesso.", numero)
        return {
            "ok":        True,
            "subint_id": subint_id,
            "numero":    numero,
            "entidade":  entidade_nome,
            "criado_em": agora,
        }

    except Exception as exc:
        logger.error("[subint] Erro ao gerar SUBINT para '%s': %s", entidade_nome, exc, exc_info=True)
        return {"ok": False, "motivo": str(exc)}


def listar_subints(limite: int = 50) -> list[dict]:
    """Lista SUBINTs gerados (sem os bytes de arquivo)."""
    try:
        with _conn() as con:
            rows = con.execute(
                "SELECT id, numero, entidade, origem, operador, criado_em "
                "FROM subint_docs ORDER BY criado_em DESC LIMIT ?",
                (limite,),
            ).fetchall()
        return [dict(r) for r in rows]
    except Exception as exc:
        logger.error("[subint] Erro ao listar: %s", exc)
        return []


def obter_bytes(subint_id: str, formato: str) -> Optional[bytes]:
    """Retorna os bytes binários do PDF ou DOCX para um SUBINT."""
    col = "pdf_bytes" if formato == "pdf" else "docx_bytes"
    try:
        with _conn() as con:
            row = con.execute(
                f"SELECT {col} FROM subint_docs WHERE id = ?",
                (subint_id,),
            ).fetchone()
        return bytes(row[0]) if row and row[0] else None
    except Exception as exc:
        logger.error("[subint] Erro ao obter bytes (%s) de %s: %s", formato, subint_id, exc)
        return None


def obter_texto(subint_id: str) -> Optional[dict]:
    """Retorna os metadados + texto de um SUBINT (sem bytes de arquivo)."""
    try:
        with _conn() as con:
            row = con.execute(
                "SELECT id, numero, entidade, origem, operador, texto, criado_em "
                "FROM subint_docs WHERE id = ?",
                (subint_id,),
            ).fetchone()
        return dict(row) if row else None
    except Exception as exc:
        logger.error("[subint] Erro ao obter texto de %s: %s", subint_id, exc)
        return None
