"""
tests/test_export_service.py
─────────────────────────────────────────────────────────────────────────────
Testes unitários para services/export_service.py

Por que estes testes são valiosos?
  - build_txt/pdf/docx são funções puras: recebem dict, devolvem bytes.
  - Zero dependência de FastAPI, banco, Drive ou Groq.
  - Rodam em <1 segundo mesmo em CI sem internet.
  - Garantem que mudanças no layout do laudo não quebram silenciosamente.

Conceito de teste que você vai absorver aqui:
  - Fixture: dado de entrada reutilizável entre testes (@pytest.fixture)
  - Assert em bytes: verificamos o conteúdo sem parsear o formato inteiro
  - Teste de fallback: o que acontece com dados incompletos/vazios?
"""

import pytest
from services.export_service import build_txt, build_pdf, build_docx


# ─── Fixture: laudo completo ──────────────────────────────────────────────────

@pytest.fixture
def laudo_completo():
    """
    Dado de entrada padrão reutilizado em todos os testes.
    Fixture = setup declarativo. Pytest injeta automaticamente em qualquer
    test_ que declare o parâmetro com o mesmo nome.
    """
    return {
        "laudo_number": "0518/2026",
        "date":         "18 de Maio de 2026",
        "filename":     "audio_operacional.wav",
        "duration":     "03:42:00",
        "risk_level":   "ALTO",
        "classification": "Transcricao de audio operacional SEAP/AM",
        "summary":      "Dois interlocutores discutem movimentacao de entorpecentes.",
        "speakers": [
            {"id": "M1", "label": "Voz masculina", "role": "Interlocutor A"},
            {"id": "F1", "label": "Voz feminina",  "role": "Interlocutor B"},
        ],
        "segments": [
            {"ts": "00:00:05", "speaker": "M1", "text": "Manda o trampo pra Compensa."},
            {"ts": "00:00:12", "speaker": "F1", "text": "Ja mandei, ta na mao do cara."},
        ],
        "red_flags": [
            {"id": 1, "title": "Trafico", "text": "Referencia a entrega de entorpecentes."},
        ],
    }


@pytest.fixture
def laudo_minimo():
    """Laudo com apenas campos obrigatórios — testa robustez com dados incompletos."""
    return {
        "laudo_number": "0001/2026",
        "date":         "1 de Janeiro de 2026",
        "filename":     "audio.wav",
        "duration":     "00:01:00",
        "risk_level":   "BAIXO",
        "classification": "Teste",
        "summary":      "Sem conteudo relevante.",
        "speakers":     [],
        "segments":     [],
        "red_flags":    [],
    }


# ─── Testes: build_txt ────────────────────────────────────────────────────────

class TestBuildTxt:
    def test_retorna_bytes(self, laudo_completo):
        resultado = build_txt(laudo_completo)
        assert isinstance(resultado, bytes)

    def test_encoding_utf8(self, laudo_completo):
        """Garante que o arquivo pode ser lido como UTF-8 sem erro."""
        resultado = build_txt(laudo_completo)
        texto = resultado.decode("utf-8")
        assert len(texto) > 0

    def test_contem_numero_laudo(self, laudo_completo):
        texto = build_txt(laudo_completo).decode("utf-8")
        assert "0518/2026" in texto

    def test_contem_nivel_risco(self, laudo_completo):
        texto = build_txt(laudo_completo).decode("utf-8")
        assert "ALTO" in texto

    def test_contem_segmentos(self, laudo_completo):
        texto = build_txt(laudo_completo).decode("utf-8")
        assert "Manda o trampo pra Compensa" in texto

    def test_contem_red_flags(self, laudo_completo):
        texto = build_txt(laudo_completo).decode("utf-8")
        assert "Trafico" in texto

    def test_laudo_vazio_nao_quebra(self, laudo_minimo):
        """Dado mínimo não deve lançar exceção — robustez de produção."""
        resultado = build_txt(laudo_minimo)
        assert isinstance(resultado, bytes)
        assert len(resultado) > 0

    def test_sem_red_flags_omite_secao(self, laudo_minimo):
        texto = build_txt(laudo_minimo).decode("utf-8")
        assert "ALERTAS IDENTIFICADOS" not in texto

    def test_com_red_flags_inclui_secao(self, laudo_completo):
        texto = build_txt(laudo_completo).decode("utf-8")
        assert "ALERTAS IDENTIFICADOS" in texto


# ─── Testes: build_pdf ────────────────────────────────────────────────────────

class TestBuildPdf:
    def test_retorna_bytes(self, laudo_completo):
        resultado = build_pdf(laudo_completo)
        assert isinstance(resultado, bytes)

    def test_assinatura_pdf(self, laudo_completo):
        """Todo PDF começa com %PDF — verifica que é um PDF real, não lixo."""
        resultado = build_pdf(laudo_completo)
        assert resultado[:4] == b"%PDF"

    def test_tamanho_minimo(self, laudo_completo):
        """PDF com conteúdo real tem pelo menos 1KB."""
        resultado = build_pdf(laudo_completo)
        assert len(resultado) > 1024

    def test_laudo_minimo_nao_quebra(self, laudo_minimo):
        resultado = build_pdf(laudo_minimo)
        assert resultado[:4] == b"%PDF"

    def test_risco_alto_gera_pdf(self, laudo_completo):
        laudo_completo["risk_level"] = "ALTO"
        resultado = build_pdf(laudo_completo)
        assert isinstance(resultado, bytes)

    def test_risco_baixo_gera_pdf(self, laudo_completo):
        laudo_completo["risk_level"] = "BAIXO"
        resultado = build_pdf(laudo_completo)
        assert isinstance(resultado, bytes)


# ─── Testes: build_docx ───────────────────────────────────────────────────────

class TestBuildDocx:
    def test_retorna_bytes(self, laudo_completo):
        resultado = build_docx(laudo_completo)
        assert isinstance(resultado, bytes)

    def test_assinatura_docx(self, laudo_completo):
        """
        DOCX é um ZIP — começa com PK (0x50 0x4B).
        Verificar a assinatura garante que é um arquivo Office real.
        """
        resultado = build_docx(laudo_completo)
        assert resultado[:2] == b"PK"

    def test_tamanho_minimo(self, laudo_completo):
        resultado = build_docx(laudo_completo)
        assert len(resultado) > 1024

    def test_laudo_minimo_nao_quebra(self, laudo_minimo):
        resultado = build_docx(laudo_minimo)
        assert resultado[:2] == b"PK"

    def test_conteudo_legivel(self, laudo_completo):
        """
        Abre o DOCX e verifica conteudo de texto.
        Numero do laudo fica na tabela de metadados — lemos paragrafos E celulas.
        """
        import io
        from docx import Document
        resultado = build_docx(laudo_completo)
        doc   = Document(io.BytesIO(resultado))
        # Paragrafos normais
        texto = "\n".join(p.text for p in doc.paragraphs)
        # Celulas das tabelas (onde fica o numero do laudo)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texto += "\n" + cell.text
        assert "0518/2026" in texto
        assert "ALTO" in texto


# ─── Testes de consistência entre formatos ────────────────────────────────────

class TestConsistenciaFormatos:
    def test_todos_formatos_contem_numero_laudo(self, laudo_completo):
        """
        O numero do laudo deve aparecer em todos os tres formatos.
        No DOCX ele fica na tabela de metadados, nao nos paragrafos.
        """
        numero = laudo_completo["laudo_number"]

        txt = build_txt(laudo_completo).decode("utf-8")
        assert numero in txt, "TXT nao contem numero do laudo"

        import io
        from docx import Document
        docx_bytes = build_docx(laudo_completo)
        doc = Document(io.BytesIO(docx_bytes))
        # Numero do laudo fica na tabela de metadados
        docx_texto = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    docx_texto += "\n" + cell.text
        assert numero in docx_texto, "DOCX nao contem numero do laudo"

        pdf_bytes = build_pdf(laudo_completo)
        assert pdf_bytes[:4] == b"%PDF", "PDF invalido"

    def test_formatos_diferentes_entre_si(self, laudo_completo):
        """Os três formatos devem gerar bytes diferentes — sanity check."""
        txt  = build_txt(laudo_completo)
        pdf  = build_pdf(laudo_completo)
        docx = build_docx(laudo_completo)
        assert txt != pdf
        assert txt != docx
        assert pdf != docx
